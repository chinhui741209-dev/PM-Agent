"""把不同格式的輸入解析成正規化的需求文字。

支援：純文字 / Markdown、Word(.docx)、PDF、Excel(.xlsx/.xls)。
Excel 額外嘗試把每列視為一筆結構化的交付物/工作項。
"""
from __future__ import annotations

import io


def parse_text(text: str) -> tuple[str, list[str]]:
    return text.strip(), []


def parse_docx(data: bytes) -> tuple[str, list[str]]:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格內容也納入
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip(), []


def parse_pdf(data: bytes) -> tuple[str, list[str]]:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(txt)
    return "\n".join(parts).strip(), []


def parse_xlsx(data: bytes) -> tuple[str, list[str]]:
    """讀第一個工作表。把每列轉成一行文字；若有看似標題的第一列，
    其餘各列的第一個非空欄位收進 deliverables。"""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    ws = wb.active
    rows: list[list[str]] = []
    for row in ws.iter_rows(values_only=True):
        cells = ["" if v is None else str(v).strip() for v in row]
        if any(cells):
            rows.append(cells)

    lines: list[str] = []
    deliverables: list[str] = []
    for i, cells in enumerate(rows):
        non_empty = [c for c in cells if c]
        lines.append(" | ".join(non_empty))
        if i > 0 and non_empty:  # 跳過標題列，把每列第一欄當交付物名稱
            deliverables.append(non_empty[0])
    return "\n".join(lines).strip(), deliverables


def parse_upload(filename: str, data: bytes) -> tuple[str, str, list[str]]:
    """回傳 (requirement_text, source_label, extracted_deliverables)。"""
    name = (filename or "").lower()
    if name.endswith(".docx"):
        text, deliv = parse_docx(data)
        return text, "docx", deliv
    if name.endswith(".pdf"):
        text, deliv = parse_pdf(data)
        return text, "pdf", deliv
    if name.endswith((".xlsx", ".xlsm", ".xls")):
        text, deliv = parse_xlsx(data)
        return text, "xlsx", deliv
    if name.endswith((".md", ".markdown")):
        return data.decode("utf-8", errors="replace").strip(), "md", []
    # 其餘當純文字
    return data.decode("utf-8", errors="replace").strip(), "txt", []
